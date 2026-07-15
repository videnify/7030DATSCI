"""Synchronise the live Word artefacts with the validated FES v1.1 results.

This script deliberately leaves the dated collection report unchanged because it is
declared as a legacy artefact in reports/README.md.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shape import InlineShape
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
BACKUP = ROOT / "reports/dissertation/archive/pre_fes_v1_1_doc_sync_2026-07-14"
MAIN = ROOT / "reports/dissertation/7030DATSCI-14_07_2026.docx"
AI = ROOT / "reports/dissertation/7030DATSCI-14_07_2026_Generative_AI_Statement.docx"
SUMMARY = ROOT / "docs/Project_Summary.docx"


def backup_once(path: Path) -> None:
    BACKUP.mkdir(parents=True, exist_ok=True)
    destination = BACKUP / path.name
    if not destination.exists():
        shutil.copy2(path, destination)


def replace_paragraphs(doc: Document, replacements: dict[int, str]) -> None:
    for index, text in replacements.items():
        if index >= len(doc.paragraphs):
            raise IndexError(f"Paragraph {index} missing from document")
        doc.paragraphs[index].text = text


def fill_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, rows):
        if len(values) != len(row.cells):
            raise ValueError(f"Expected {len(row.cells)} cells, got {len(values)}")
        for cell, value in zip(row.cells, values):
            cell.text = value


def update_cached_toc(doc: Document) -> None:
    """Update the visible cached TOC values without removing Word field codes."""
    paragraphs = [
        paragraph
        for paragraph in doc._element.xpath(".//w:sdt//w:p")
        if paragraph.xpath(".//w:t")
    ]
    pages = [
        2, 3, 3, 3, 3, 4, 5, 5, 6, 6, 7, 7, 7, 7, 8, 8, 8, 9, 9, 9,
        10, 11, 12, 13, 15, 15, 16, 18, 18, 21, 22, 23, 24, 24, 26, 26, 27,
        28, 28, 30,
    ]
    if len(paragraphs) != len(pages):
        raise ValueError(f"Unexpected TOC size: {len(paragraphs)}")
    heading_names = {
        0: "Table of Contents",
        1: "List of Figures",
        2: "List of Tables",
        3: "List of Appendix Figures",
        4: "List of Appendix Tables",
    }
    for index, (paragraph, page) in enumerate(zip(paragraphs, pages)):
        texts = paragraph.xpath(".//w:t")
        if not texts:
            continue
        texts[-1].text = str(page)
        if index in heading_names:
            texts[0].text = heading_names[index]


def replace_inline_image(doc: Document, paragraph_index: int, image_path: Path, width: float) -> None:
    paragraph = doc.paragraphs[paragraph_index]
    blips = paragraph._p.xpath(".//a:blip")
    inlines = paragraph._p.xpath(".//wp:inline")
    if len(blips) != 1 or len(inlines) != 1:
        raise ValueError(f"Expected one inline image in paragraph {paragraph_index}")
    relationship_id = blips[0].get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    image_part = doc.part.related_parts[relationship_id]
    image_part._blob = image_path.read_bytes()
    # Word may retain an SVG preference alongside the PNG fallback. LibreOffice
    # renders that preferred SVG, so remove the stale extension when replacing.
    for extension_list in blips[0].xpath("./a:extLst"):
        blips[0].remove(extension_list)
    with Image.open(image_path) as image:
        height = width * image.height / image.width
    shape = InlineShape(inlines[0])
    shape.width = Inches(width)
    shape.height = Inches(height)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def sync_dissertation() -> None:
    backup_once(MAIN)
    doc = Document(MAIN)
    replace_paragraphs(
        doc,
        {
            4: "How do political and macroeconomic events affect the US stock market, and can these events improve short-term market prediction?",
            18: "List of Figures",
            19: "Figure 1. Data lineage from current raw sources to FES v1.1.\t13",
            20: "Figure 2. Current predictive modelling flow.\t14",
            21: "Figure 3. Event landscape overview.\t15",
            22: "Figure 4. Current event-study and causal-evidence panel.\t17",
            23: "Figure 5. Predictive pipeline overview.\t20",
            24: "Figure 6. Integrated test-period dashboard (2023–2025).\t22",
            26: "List of Tables",
            27: "Table 1. Summary of data sources.\t7",
            28: "Table 2. FinBERT scoring parameters.\t9",
            29: "Table 3. Statistical tests applied to event-study CAR.\t10",
            30: "Table 4. Structural causal model edges.\t11",
            31: "Table 5. Refutation tests for the pooled DoWhy estimate.\t11",
            32: "Table 6. FES v1.1 feature groups.\t12",
            33: "Table 7. XGBoost and LightGBM hyperparameters.\t14",
            34: "Table 8. Mean CAR by event type.\t16",
            35: "Table 9. DoWhy backdoor estimates.\t18",
            36: "Table 10. Current Random Forest top-ten importance.\t19",
            37: "Table 11. Held-out model performance.\t21",
            39: "List of Appendix Figures",
            40: "Figure A1. Current CAR distribution by event type.\t28",
            41: "Figure A2. Current DoWhy causal estimates with 95% confidence intervals.\t28",
            42: "Figure A3. FES v1.1 feature-matrix validation summary.\t28",
            43: "Figure A4. Baseline validation summary.\t29",
            44: "Figure A5. Model-evaluation validation summary.\t29",
            45: "Figure A6. SPY rebased price over the study period.\t29",
            47: "List of Appendix Tables",
            48: "Table B1. Summary of validated outputs by phase.\t30",
            52: "This dissertation evaluates whether public communications and macroeconomic policy events affect S&P 500 (SPY) returns and whether event-derived information improves next-day prediction. It integrates an event study, DoWhy causal inference, feature engineering, explainable machine learning and a market-only baseline within one reproducible eight-notebook pipeline.",
            53: "The current event catalogue contains 1,005 observations: 916 American Presidency Project (APP) documents scored with FinBERT and 89 structured FOMC decisions. A 344-row high-impact subset supports the event-study stage, which yields 264 usable event observations across monetary, geopolitical, regulatory, trade and energy categories. The frozen FES v1.1 matrix contains 92 features and 2,477 observations, split chronologically into 1,727 training and 750 test rows.",
            54: "No event-type mean-CAR null is rejected after Benjamini–Hochberg FDR correction (minimum raw p = 0.1162; minimum q = 0.5810); all five 95% intervals cross zero and the largest |Cohen's d| is 0.239. The separate pooled DoWhy estimate is +0.005601 in next-day log-return units (95% CI [+0.002295, +0.008907], p = 0.0009), conditional on the stated DAG and no-unmeasured-confounding assumption.",
            55: "Random Forest importance is dominated by market and macroeconomic features; no event feature appears in its top ten, so the RQ2 null is not rejected. On the common test split, Baseline_LASSO and Event_LASSO are identical constant predictors (RMSE 0.009631; directional accuracy 57.47%). XGBoost and LightGBM are worse on RMSE and directional accuracy, and neither clears the Bonferroni-corrected comparison protocol. The RQ3 null is therefore not rejected.",
            56: "The results show a difference between explanatory and predictive evidence: a pooled conditional causal association is detectable, yet event-enhanced models do not improve out-of-sample forecasting. The project reports this null predictive result directly and preserves validation hashes, contracts and known exceptions so that the conclusions remain auditable.",
            83: "A documented sentiment pipeline in which 916 APP titles are scored by FinBERT and 89 FOMC decisions retain a structured policy signal; the lexicon implementation is retained only as fallback and history.",
            84: "A reproducible feature-importance analysis using Random Forest and SHAP diagnostics, with explicit separation between descriptive importance and formal hypothesis testing.",
            92: "Five data sources were assembled for January 2015 to December 2025: SPY/VIX market data, APP documents, FRED macroeconomic indicators, FOMC decisions and GDELT daily risk history. GDELT is retained as a candidate contextual control but is deliberately excluded from the frozen causal DAG and FES v1.1. Table 1 summarises the source boundary.",
            93: "Table 1. Summary of data sources.",
            97: "The SPY price history comprises 2,765 trading days. Log returns define the next-day modelling target because they are time-additive; the event study separately uses simple percentage returns in its constant-mean abnormal-return calculation. This distinction is maintained throughout the pipeline.",
            99: "Presidential communications were sourced from the American Presidency Project (APP), a non-partisan archive maintained by UC Santa Barbara. The current raw APP collection contains 10,892 dated records; the reproducible analytical input is the 916-row economic-policy subset persisted as app_presidential_documents_economic.parquet.",
            100: "Document titles, rather than full text, are the NLP input. This supports reproducible batch scoring but can miss nuance in longer communications and contributes to the high neutral share; full-text scoring is therefore treated as future work rather than implied by the current results.",
            106: "The main limitations are title-only sentiment, clustered events, observational causal identification and strong residual non-normality. BH-FDR controls the five event-type family but does not repair temporal dependence within clustered communications. GDELT covers 4,018 daily observations but remains outside the current DoWhy DAG, feature matrix and reported RQ evidence. Results are limited to SPY and a single chronological test period.",
            113: "The current unified catalogue combines 916 FinBERT-scored APP economic-policy documents with 89 structured FOMC decisions. Rule-based taxonomy assignment produces eight event categories, while 344 records meet the high-impact definition used to seed the event-study stage.",
            115: "Table 2. FinBERT scoring parameters fixed in the project configuration.",
            117: "The curated lexicon scorer is retained only as a fallback when FinBERT is unavailable and as a historical prototype. It is not the current primary APP treatment. A legacy persisted value, sentiment_method = lexicon, remains in causal_estimates.parquet as an accepted known exception; the current treatment actually combines APP FinBERT and structured FOMC sentiment.",
            118: "The full-history GDELT daily file contains 4,018 risk observations. It remains available in the master dataset as a candidate contextual control, but is intentionally excluded from the frozen DoWhy graph and FES v1.1; no current RQ result relies on GDELT.",
            119: "Daily sentiment aggregates contain 739 dated rows and 18 columns, including explicit event-occurrence counts and GDELT context fields. Days without classified APP/FOMC events are not silently interpreted as neutral communications.",
            125: "After date/category alignment and event-window availability checks, the event study contains 264 usable observations across monetary (45), geopolitical (60), regulatory (50), trade (62) and energy (47) categories. It uses a [−5,+10] event window and the frozen pre-event estimation rule.",
            127: "Table 3. Statistical tests applied to event-study CAR.",
            129: "Event clustering can violate the independence assumption of the one-sample event-type tests. The five-type family is now reported with 95% t intervals, Cohen's d and Benjamini–Hochberg FDR-adjusted q-values. This controls multiplicity across types but does not correct within-type temporal dependence, which remains a limitation.",
            130: "car_results.parquet contains 264 rows with CAR, mean abnormal return, window counts, per-event t-statistic and significance flag, VIX and the aligned sentiment aggregate. Its current schema and file hash are validated downstream.",
            133: "Table 4. Structural causal model edges in the primary DoWhy estimate.",
            135: "VIX regime is a binary indicator equal to one when VIX exceeds 20. The current causal treatment is the combined same-day APP FinBERT and structured FOMC sentiment aggregate. The persisted sentiment_method value lexicon is a legacy label only and is documented as an accepted known exception.",
            141: "Random-common-cause, placebo-treatment and data-subset refuters were applied to the pooled estimate. These checks support stability against those specific perturbations but do not prove the absence of unmeasured confounding and were not repeated for every category-specific estimate.",
            139: "Table 5. Refutation tests applied to the primary DoWhy estimate.",
            142: "Five category estimates are saved in causal_estimates.parquet. Geopolitical and energy each have only one non-zero treatment day, so their narrow intervals must not be interpreted as robust category-level evidence. The pooled estimate is the principal reported causal result.",
            144: "Feature engineering consumes the validated master dataset and CAR output. FES v1.1 contains 92 predictors across six categories: market (27), macro (16), sentiment (23), event (14), temporal (5) and interaction (7).",
            145: "Table 6. Feature groups comprising the frozen 92-feature matrix (FES v1.1).",
            147: "All 92 contracted features are retained for event-enhanced models. Three zero-training-variance features from FES v1.0 were removed explicitly in v1.1; leakage, source, split and schema checks pass in feature_matrix_validation.json.",
            151: "After deterministic trimming, the matrix contains 2,477 rows: 1,727 training observations and 750 held-out test observations. The chronological split is shared by every predictive model.",
            152: "The 27 market features are the only predictors available to Baseline_LASSO. The remaining 65 features are available only to event-enhanced candidates, enforcing the baseline boundary at model consumption.",
            149: "Figure 1. Data lineage from current raw sources to the validated FES v1.1 feature matrix and downstream evidence.",
            154: "Four models address RQ3: market-only Baseline_LASSO and three 92-feature candidates—Event_LASSO, XGBoost and LightGBM. All are compared on the same 750-row held-out period.",
            156: "Figure 2. Current predictive modelling flow under FES v1.1 / MCP v1.0.",
            158: "Baseline_LASSO and Event_LASSO use the same L1-regularised linear estimator with different feature scopes. Baseline_LASSO reads only 27 market features; Event_LASSO reads all 92 FES v1.1 features. Cross-validation drives every coefficient in both models to zero, so each predicts the training-split mean on every held-out day.",
            162: "Table 7. Selected hyperparameters for XGBoost and LightGBM.",
            159: "XGBoost and LightGBM were trained as the two non-linear event-enhanced candidates on FES v1.1. Their strong training scores and negative held-out R² values make overfitting a central diagnostic finding rather than evidence of predictive improvement.",
            165: "Explainability is reported as diagnostic evidence. Random Forest supplies the reproducible RQ2 ranking and SHAP summarises fitted candidate behaviour; Event_LASSO has zero non-zero coefficients under the current matrix, so its SHAP contribution is identically zero.",
            166: "",
            169: "The current catalogue contains 1,005 events: 916 APP economic-policy documents and 89 FOMC decisions. It includes 344 high-impact records. The catalogue, rather than the earlier 11,664-row prototype snapshot, is the authoritative input to the current downstream pipeline.",
            170: "Sentiment labels are neutral for 929 events (92.4%), positive for 43 (4.3%) and negative for 33 (3.3%). APP records use FinBERT and FOMC records use the structured decision signal; no current event is sourced from the lexicon fallback.",
            172: "Figure 3. Event landscape overview.",
            178: "The usable event-study output contains 264 observations across five categories. Table 8 reports each mean CAR with its 95% t interval, sample size, raw p-value, BH-FDR q-value and Cohen's d; Figure 4 presents the current distributions and causal estimates.",
            179: "Table 8. Mean cumulative abnormal return by event type.",
            181: "No event-type null is rejected after BH-FDR correction: the minimum q-value is 0.5810 for monetary events and all five 95% intervals cross zero. Standardised effects are negligible to small; monetary has the largest magnitude (d = −0.239). Fifteen of 264 rows carry an individual-window diagnostic flag, but those flags do not establish an event-type mean effect.",
            183: "Figure 4. Current event-study and causal-evidence panel.",
            187: "The pooled DoWhy backdoor model estimates that a one-unit increase in the combined same-day APP + FOMC sentiment treatment is associated with +0.005601 in next-day SPY log return (95% CI [+0.002295, +0.008907], p = 0.0009), conditional on the frozen VIX-regime and prior-return adjustment set. This is the headline causal estimate, not proof of an unqualified causal effect.",
            188: "Table 9. DoWhy backdoor estimates for sentiment and next-day SPY return.",
            190: "Category estimates are: monetary +0.007094, trade −0.001337, geopolitical −0.005918, regulatory +0.003617 and energy −0.002516. Trade is the only interval that spans zero, but geopolitical and energy each have only one non-zero treatment observation; these sparse results are reported as fragile diagnostics, not strong category claims.",
            191: "The pooled refuters do not materially overturn the headline estimate under the implemented checks. They do not test every category estimate and cannot rule out omitted confounding, functional-form error or treatment-measurement error; those limits remain part of the interpretation.",
            193: "FES v1.1 contains 92 validated features. The current Random Forest importance artefact is reproducible from Notebook 07 and is used for RQ2 descriptive ranking; it is no longer treated as an orphaned legacy model.",
            195: "Table 10. Top ten features by current Random Forest importance.",
            197: "The current top ten are entirely market or macro features. log_return_hi ranks first (0.08345), followed by return_lag1d (0.08114); the highest event feature is outside the top ten. Accordingly, the RQ2 null—that event-derived features do not dominate predictive importance—is not rejected.",
            200: "The current predictive panel reports category and feature-level Random Forest importance, XGBoost SHAP diagnostics, out-of-sample changes versus baseline, predicted-versus-actual returns and rolling information coefficient. XGBoost is displayed because it has the highest defined test IC, not because it wins the model comparison.",
            199: "Figure 5. Predictive pipeline overview under FES v1.1 / MCP v1.0.",
            203: "Table 11 reports held-out performance for the four models on the common 750-row next-day SPY return target.",
            204: "Table 11. Model performance on the 750-row held-out test set.",
            206: "Baseline_LASSO and Event_LASSO produce identical constant predictions (RMSE 0.009631; MAE 0.006548; R² −0.001506; directional accuracy 57.47%). XGBoost records RMSE 0.009656 and directional accuracy 48.93%; LightGBM records 0.009700 and 44.27%. Diebold–Mariano p-values are 0.610 for XGBoost and 0.672 for LightGBM; the Event_LASSO loss differential is identically zero. No candidate clears the Bonferroni α = 0.0167 promotion rule.",
            207: "XGBoost has the highest defined test information coefficient (0.02495), followed by LightGBM (−0.05058); IC is undefined for the two constant LASSO predictions. XGBoost falls from training R² 0.2225 and directional accuracy 65.55% to test R² −0.0067 and 48.93%, providing direct evidence of overfitting.",
            208: "Figure 5 therefore uses XGBoost only as the most informative non-constant diagnostic. Its rolling 63-day information coefficient varies materially through time and does not establish stable predictive skill.",
            211: "The integrated dashboard shows rebased SPY, VIX, 21-session event counts, the XGBoost prediction diagnostic and an illustrative sign-following series. The last two panels are diagnostics only, not a winner declaration, tradeable strategy or net-of-costs backtest.",
            210: "Figure 6. Integrated test-period dashboard (2023–2025).",
            214: "Random Forest importance assigns 71.9% cumulatively to market features and 18.2% to macro features; event features receive 3.9%. XGBoost's SHAP diagnostic is led by unemployment and return_lag21d. These fitted-model summaries support interpretation but do not alter the failed out-of-sample promotion test.",
            215: "Figure 5 presents current feature and SHAP summaries directly from the validated FES v1.1 / MCP v1.0 artefacts. No historical Event_LASSO beeswarm is used in the current evidential chain because Event_LASSO is a zero-coefficient constant model.",
            216: "Temporal variation in XGBoost's rolling IC is shown descriptively. A formal sub-period stability test was not pre-specified and is not claimed.",
            218: "Residuals are heavy-tailed for every model (Jarque–Bera p approximately zero). Durbin–Watson values are 2.112 for the constant LASSO models, 1.989 for XGBoost and 2.032 for LightGBM. Persisted block-bootstrap RMSE intervals and Wilson directional-accuracy intervals are therefore preferred to normal-theory certainty claims.",
            222: "The event study and causal model answer different questions. All five event-type mean-CAR tests remain null after BH-FDR correction, with intervals crossing zero and only negligible-to-small standardised effects, whereas the pooled adjusted sentiment estimate is positive. The latter is conditional on the specified DAG, linear estimator and no-unmeasured-confounding assumption and should not be treated as experimental proof.",
            223: "Category-specific causal estimates do not provide a clean alternative story: geopolitical and energy are driven by one non-zero treatment observation each, while trade spans zero. The final RQ1 conclusion is qualified: the pooled conditional association is detectable, but multiplicity-controlled event-type abnormal-return evidence is null.",
            225: "The pooled magnitude, approximately +0.56% for a full-unit treatment change, is economically non-trivial. However, treatment values are continuous and rarely span a full unit, and refuters were applied only to the pooled model. Economic interpretation must therefore remain proportional to realistic treatment changes.",
            226: "Only 15 of 264 event-study rows carry an individual-window diagnostic flag. Those observations do not substitute for the five event-type mean tests, for which BH-FDR rejects 0/5 nulls.",
            229: "For RQ2, the current evidence does not show event features dominating prediction. The Random Forest top ten contains only market and macro variables; event features account for 3.9% of category-level importance. XGBoost SHAP likewise highlights macro and market predictors.",
            230: "For RQ3, no event-enhanced model beats the market-only baseline on the required out-of-sample criteria. Event_LASSO collapses to the same constant as Baseline_LASSO, while both tree models are worse on RMSE and directional accuracy. This is a clear null result rather than a near-win.",
            232: "Three mechanisms are plausible: daily return predictability is intrinsically weak; sparse event information may already be reflected in price and volatility variables; and the flexible tree models overfit the training period. The 750-row test window also limits power, but extending it would be future validation rather than grounds to reinterpret the current result.",
            234: "The baseline's 57.47% directional accuracy is a base-rate artefact of its constant positive prediction, not evidence of genuine timing skill. RMSE, information coefficient, paired tests and interval estimates are therefore essential to the interpretation.",
            237: "Limitations include title-only APP sentiment, a 92.4% neutral label share, clustered events, observational identification, heavy-tailed residuals, and a single asset and test period. The RQ1 multiplicity/effect-size table is complete; its BH-FDR correction does not remove the event-dependence limitation. The legacy causal field label is documented rather than silently rewritten.",
            239: "GDELT provides 4,018 daily observations but is excluded from the current causal DAG and FES v1.1, so no reported result depends on it. Future work may evaluate it as a pre-specified control, but adding it now would change the frozen design.",
            240: "Random Forest importance is now reproducible in Notebook 07. Its remaining limitation is interpretive: impurity importance can favour continuous variables, so SHAP is used as corroborating diagnostic evidence rather than a second formal hypothesis test.",
            243: "The remaining dissertation tasks are citation, number-to-source and cross-reference verification followed by final proofreading. Analytical extensions include cluster-aware event-study inference, full-text sentiment scoring and external validation on other assets.",
            244: "Future causal work should pre-specify alternative adjustment sets, repeat robustness checks by category where sample support permits, and investigate treatment sparsity before interpreting narrow intervals.",
            246: "Future predictive work should diagnose XGBoost overfitting, test stability across non-overlapping periods and assets, and preserve a strictly untouched evaluation set. Any new model or feature source would require a versioned contract amendment.",
            248: "This dissertation contributes a reproducible, version-controlled pipeline linking event detection, abnormal-return analysis, causal inference, FES v1.1 feature engineering and MCP v1.0 model comparison. Its central value is the separation of explanatory evidence from out-of-sample predictive evidence.",
            249: "RQ1 receives a qualified final answer: BH-FDR rejects 0/5 event-type mean-CAR nulls (minimum q = 0.5810; maximum |d| = 0.239), while the separate pooled conditional sentiment estimate is positive.",
            250: "RQ2 does not reject its null: current Random Forest importance is dominated by market and macro variables and contains no event feature in the top ten. SHAP is treated as supporting model diagnostics.",
            251: "RQ3 does not reject its null: Event_LASSO equals the constant baseline, and XGBoost and LightGBM fail the Bonferroni-corrected promotion protocol. No event-enhanced model demonstrates superior next-day prediction.",
            252: "These results show why causal, descriptive and predictive claims must remain distinct. A conditional pooled estimate can coexist with null event-type means and null forecasting gains without contradiction.",
            253: "The principal limitations are treatment measurement, event dependence, sparse category treatments, observational causal identification, heavy-tailed residuals and restricted external validity.",
            254: "Future work should consider cluster-aware event-study inference, full-text NLP, alternative causal estimators, pre-specified GDELT controls, richer market data and external replication.",
            255: "Taken together, the project provides an auditable example of how political and macroeconomic information can be processed and tested without converting a weak or null predictive result into an overstated claim.",
            257: "The end-to-end eight-notebook pipeline now runs with current saved outputs, FES v1.1 and all model/visualisation validations passing. The most important achievement is not a positive forecast result, but an auditable workflow that exposes null findings and known exceptions.",
            258: "If starting again, I would freeze the treatment definition and multiplicity plan earlier, and treat GDELT as an optional pre-specified control from the outset. The present full-history GDELT file is useful context but remains outside the validated evidential chain.",
            259: "I would also allocate more time to a labelled validation set and full-text APP analysis. Title-only FinBERT scoring is reproducible but cannot capture all policy nuance.",
            260: "The contrast between XGBoost's training and test results reinforced the importance of chronological validation and simple baselines in financial forecasting. The final interpretation therefore privileges held-out performance over model complexity.",
            286: "Figure A1. Current CAR distribution by event type.",
            289: "Figure A2. Current DoWhy estimates with 95% confidence intervals; the persisted lexicon method value is a legacy label, not the current treatment definition.",
            291: "Figure A3. Notebook 05 validation summary for the 92-feature FES v1.1 matrix and chronological split.",
            292: "The matrix contains 2,477 rows: 1,727 training and 750 held-out test observations.",
            294: "Figure A4. Notebook 06 baseline validation summary.",
            295: "Baseline_LASSO reduces to a constant but improves on a naive persistence comparator; its 57.47% directional accuracy is a base-rate artefact.",
            298: "Figure A5. Notebook 07 model-evaluation validation summary.",
            299: "No event-enhanced candidate beats Baseline_LASSO on both Bonferroni-corrected comparison legs.",
            301: "Figure A6. SPY rebased price over the study period.",
            302: "Market context only; this is not a strategy or performance claim.",
            303: "The study period covers multiple volatility and policy regimes between 2015 and 2025.",
            305: " Appendix B: Data and Code",
            306: "The submission repository contains all eight executed notebooks, supporting Python modules, validated Parquet/JSON artefacts, architecture diagrams and publication figures. Current contracts and known exceptions are documented in the Research Bible and validation files.",
            307: "Table B1. Summary of validated outputs by phase.",
        },
    )

    fill_table(
        doc.tables[0],
        [
            ["Source", "Data Type", "Volume", "Date Range"],
            ["Yahoo Finance (yfinance)", "SPY OHLCV and VIX", "2,765 SPY trading days", "Jan 2015–Dec 2025"],
            ["American Presidency Project", "Raw / economic-policy documents", "10,892 raw; 916 analytical", "Jan 2015–Dec 2025"],
            ["Federal Reserve (FRED)", "Macroeconomic indicators", "Validated monthly series", "Jan 2015–Dec 2025"],
            ["Federal Reserve (FOMC)", "Rate decisions", "89 meetings", "Jan 2015–Dec 2025"],
            ["GDELT Project", "Candidate geopolitical-risk control", "4,018 daily rows", "2015–2025; excluded from DAG/FES"],
        ],
    )
    fill_table(
        doc.tables[5],
        [
            ["Category", "Count", "Contents"],
            ["Market", "27", "Returns, lags, cumulative returns, volatility, momentum and technical indicators"],
            ["Macroeconomic and VIX", "16", "VIX, rates, inflation, unemployment, yields and regime/change indicators"],
            ["Sentiment", "23", "Overall and event-category aggregates, lags, rolling summaries and polarity flags"],
            ["Event", "14", "CAR-derived values, event counts, recency and event-day flags"],
            ["Temporal", "5", "Cyclical calendar encodings and quarter"],
            ["Interaction", "7", "Pre-specified sentiment, event, volatility and momentum interactions"],
        ],
    )
    fill_table(
        doc.tables[7],
        [
            ["Event Type", "Mean CAR [95% CI]", "N", "Raw p / BH q", "Cohen's d"],
            ["Monetary", "−0.01431 [−0.03231,+0.00369]", "45", "0.1162 / 0.5810", "−0.239"],
            ["Geopolitical", "−0.00510 [−0.01472,+0.00452]", "60", "0.2928 / 0.7320", "−0.137"],
            ["Regulatory", "−0.00205 [−0.01539,+0.01128]", "50", "0.7584 / 0.7584", "−0.044"],
            ["Trade", "−0.00188 [−0.01304,+0.00928]", "62", "0.7377 / 0.7584", "−0.043"],
            ["Energy", "+0.00346 [−0.00837,+0.01529]", "47", "0.5589 / 0.7584", "+0.086"],
        ],
    )
    rq1_table = doc.tables[7]
    rq1_table.autofit = False
    rq1_widths = [Inches(1.25), Inches(2.25), Inches(0.45), Inches(1.35), Inches(0.85)]
    for row in rq1_table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, rq1_widths)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    fill_table(
        doc.tables[8],
        [
            ["Estimate", "Causal Effect", "CI Lower (95%)", "CI Upper (95%)", "Interpretation"],
            ["Pooled", "+0.005601", "+0.002295", "+0.008907", "Headline conditional estimate"],
            ["Monetary", "+0.007094", "+0.002577", "+0.010952", "CI excludes zero"],
            ["Trade", "−0.001337", "−0.004427", "+0.003225", "CI spans zero"],
            ["Geopolitical", "−0.005918", "−0.011837", "−0.003865", "Sparse: one non-zero day"],
            ["Regulatory", "+0.003617", "+0.000470", "+0.007080", "CI excludes zero"],
            ["Energy", "−0.002516", "−0.005031", "−0.002278", "Sparse: one non-zero day"],
        ],
    )
    fill_table(
        doc.tables[9],
        [
            ["Feature", "Importance", "Group", "Description"],
            ["log_return_hi", "0.08345", "Market", "High-to-open return"],
            ["return_lag1d", "0.08114", "Market", "One-day lagged return"],
            ["return_lag3d", "0.04122", "Market", "Three-day lagged return"],
            ["return_lag5d", "0.04106", "Market", "Five-day lagged return"],
            ["vix", "0.03852", "Macro", "VIX level"],
            ["cum_return_5d", "0.03473", "Market", "Five-day cumulative return"],
            ["vix_change_5d", "0.03301", "Macro", "Five-day VIX change"],
            ["price_vs_ma200", "0.03249", "Market", "Price relative to 200-day moving average"],
            ["momentum_63d", "0.03222", "Market", "63-day momentum"],
            ["vix_change_1d", "0.03216", "Macro", "One-day VIX change"],
        ],
    )
    fill_table(
        doc.tables[10],
        [
            ["Model", "RMSE", "MAE", "R²", "Dir Acc", "IC", "Promotion"],
            ["Baseline_LASSO", "0.009631", "0.006548", "−0.001506", "57.47%", "Undefined", "Reference"],
            ["Event_LASSO", "0.009631", "0.006548", "−0.001506", "57.47%", "Undefined", "No"],
            ["XGBoost", "0.009656", "0.006632", "−0.006713", "48.93%", "0.02495", "No"],
            ["LightGBM", "0.009700", "0.006770", "−0.015825", "44.27%", "−0.05058", "No"],
        ],
    )
    fill_table(
        doc.tables[11],
        [
            ["#", "Notebook", "Key Data Outputs", "Figures"],
            ["1", "01_data_collection.ipynb", "Validated raw Parquet sources", "Current collection figures"],
            ["2", "02_eda.ipynb", "master_dataset.parquet + validation", "EDA and stationarity figures"],
            ["3", "03_event_detection.ipynb", "events_tagged.parquet; daily_sentiment.parquet", "03a–03d"],
            ["4", "04_causal_analysis.ipynb", "car_results.parquet; causal_estimates.parquet", "04a–04d"],
            ["5", "05_feature_engineering.ipynb", "FES v1.1 matrix/profile/validation", "05 learning outcome"],
            ["6", "06_model_training.ipynb", "Baseline model/predictions/validation", "06 learning outcome"],
            ["7", "07_model_evaluation.ipynb", "Comparison, tests, RF/SHAP and validation", "07 learning outcome"],
            ["8", "08_results_visualisation.ipynb", "Figure manifest/validation", "08a–08d"],
        ],
    )

    replacements = {
        148: Path("/private/tmp/7030-svg-qa/data_lineage.png"),
        155: Path("/private/tmp/7030-svg-qa/modelling_flow.png"),
        171: ROOT / "reports/figures/08a_event_landscape.png",
        182: ROOT / "reports/figures/08b_causal_evidence.png",
        198: ROOT / "reports/figures/08c_predictive_pipeline.png",
        209: ROOT / "reports/figures/08d_full_dashboard.png",
        285: ROOT / "reports/figures/04a_car_by_event_type.png",
        288: ROOT / "reports/figures/04c_causal_estimates.png",
        290: ROOT / "reports/figures/05_learning_outcome.png",
        293: ROOT / "reports/figures/06_learning_outcome.png",
        297: ROOT / "reports/figures/07_learning_outcome.png",
        300: ROOT / "reports/figures/01_spy_rebased_price.png",
    }
    for paragraph, image in replacements.items():
        replace_inline_image(doc, paragraph, image, 6.15 if paragraph not in {148, 155} else 5.8)

    update_cached_toc(doc)

    doc.core_properties.title = "Causal Event-Driven Market Impact Modelling"
    doc.core_properties.subject = "Validated dissertation draft — FES v1.1 / MCP v1.0"
    doc.core_properties.comments = "Synchronized with executed Notebooks 01–08 and the completed RQ1 reporting gate on 2026-07-15."
    doc.save(MAIN)


def sync_ai_statement() -> None:
    backup_once(AI)
    doc = Document(AI)
    replace_paragraphs(
        doc,
        {
            3: "Claude (Anthropic) was used as a coding and repository-maintenance assistant during the project. OpenAI Codex was also used during the final consistency pass. The workflow involved describing bounded analytical or documentation tasks, inspecting proposed changes, running the affected notebooks/tests, and accepting changes only after review.",
            4: "AI-assisted tasks included data-collection and API integration, event-taxonomy debugging, sentiment-pipeline support, DoWhy compatibility fixes, FES v1.1 feature-contract implementation, baseline/model validation, publication-figure refreshes, architecture diagrams, documentation synchronization and Word-document formatting.",
            5: "All executable code and persisted outputs are included in the repository. AI assistance does not replace authorship or verification: research questions, frozen design decisions, interpretation and final submission responsibility remain my own, and claims were checked against the saved Parquet/JSON validation artefacts and executed notebooks.",
            8: "Generative AI was used to help edit, condense and synchronize prose in this report and supporting documentation after the analyses were executed. It also helped identify stale numerical claims and align captions, tables and diagrams with validated outputs. I reviewed the resulting text, checked numerical statements against source artefacts and remain responsible for the final wording, citations and argument.",
            9: "Repository Maintenance, Debugging, and Documentation Assistance\n(2026-07-13 to 2026-07-15)",
            10: "In the final repository pass (13–15 July 2026), Claude and OpenAI Codex were used for bounded debugging, validation and documentation tasks. Their role included proposing edits and running mechanical checks; the researcher reviewed the diffs and retained responsibility for methodological choices and interpretation.",
            11: "• Diagnosing environment and library compatibility issues that blocked current notebook execution, then validating the repaired pipeline.",
            12: "• Correcting event-taxonomy and data-pipeline inconsistencies discovered during clean notebook re-execution.",
            13: "• Rebuilding the full-history GDELT context file while preserving the frozen design decision to exclude GDELT from the current causal DAG and FES v1.1 evidence.",
            14: "• Implementing and validating FES v1.1, the baseline boundary, model-comparison artefacts and the accepted legacy causal-field-label exception.",
            15: "• Refreshing learning-outcome and dissertation figures from the current persisted outputs.",
            16: "• Auditing and updating the six architecture SVGs, including programmatic XML validation and rendered visual inspection.",
            17: "• Synchronizing the root and folder READMEs, Research Bible, dissertation draft, project summary and AI disclosure with the executed Notebooks 01–08 and validation files.",
            18: "All AI-assisted changes were reviewed before acceptance. The underlying research findings, methodological decisions and interpretation remain my own, and I am able to explain and defend the submitted work.",
        },
    )
    doc.paragraphs[9].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.core_properties.comments = "Updated disclosure for the final AI-assisted consistency pass on 2026-07-15."
    doc.save(AI)


def sync_project_summary() -> None:
    backup_once(SUMMARY)
    doc = Document(SUMMARY)
    doc.sections[0].header.paragraphs[0].text = (
        "DATSCI7030 — Causal Event-Driven Market Impact Modelling · "
        "Version 1.2 · Ibrahim Haroun · LJMU 2025–2026"
    )
    replace_paragraphs(
        doc,
        {
            7: "DATSCI7030 evaluates whether presidential economic-policy communications and FOMC decisions affect SPY returns, and whether event-derived information improves next-day prediction beyond a market-only baseline.",
            8: "The validated pipeline combines event-study abnormal returns, DoWhy causal inference, a frozen 92-feature matrix (FES v1.1), Random Forest/SHAP diagnostics and a Bonferroni-controlled comparison of Event_LASSO, XGBoost and LightGBM against Baseline_LASSO.",
            9: "SPY is the sole target asset in the current evidential chain. QQQ, GLD and TLT appear only as contextual EDA series and are not modelled targets.",
            13: "The project is implemented in eight sequential, fully executed notebooks:",
            15: "Collect and validate SPY/VIX, APP, FRED, FOMC and GDELT source data for 2015–2025.",
            16: "Build and validate the master dataset; describe distributions and run ADF/KPSS stationarity diagnostics.",
            17: "Create the 1,005-event APP + FOMC catalogue using FinBERT for APP and structured FOMC sentiment.",
            18: "Estimate event-window CAR and a conditional DoWhy effect using the frozen treatment and adjustment set.",
            19: "Build FES v1.1: 92 features and 2,477 rows with leakage, variance, source and split checks.",
            20: "Train the 27-feature market-only Baseline_LASSO on a chronological split.",
            21: "Train/evaluate Event_LASSO, XGBoost and LightGBM; generate RF/SHAP diagnostics and paired tests.",
            22: "Produce hash-bound publication figures 08a–08d and synchronize dissertation evidence.",
            26: "For each usable high-impact event, abnormal returns are the realised simple SPY return minus the pre-event constant mean. CAR sums abnormal returns over [−5,+10] trading days; the current output contains 264 observations across five categories.",
            29: "Five two-sided event-type mean-CAR tests are reported with 95% t intervals, Cohen's d and Benjamini–Hochberg FDR correction. No null is rejected: minimum raw p=0.1162, minimum q=0.5810, every interval crosses zero and maximum |d|=0.239.",
            32: "The DoWhy model treats the combined same-day APP + FOMC sentiment aggregate as treatment and next-day SPY log return as outcome. VIX regime and prior-day return form the frozen adjustment set.",
            34: "Treatment (T): combined same-day APP FinBERT + structured FOMC sentiment",
            35: "Outcome (Y): next-day SPY log return",
            36: "Confounders (Z): VIX regime and prior-day SPY return",
            37: "Context only: GDELT daily risk is excluded from the frozen DAG and FES v1.1",
            39: "DoWhy uses backdoor linear regression with random-common-cause, placebo-treatment and data-subset refuters on the pooled estimate. The current pooled effect is +0.005601 (95% CI [+0.002295,+0.008907], p=0.0009), subject to observational-design assumptions.",
            42: "The analytical event input contains 916 APP economic-policy titles scored with FinBERT and 89 FOMC decisions scored by a structured rule. The current label distribution is 929 neutral, 43 positive and 33 negative.",
            43: "The lexicon scorer is fallback/history only. The causal output's sentiment_method=lexicon value is a documented legacy label and does not describe the current combined treatment.",
            46: "Four models share the same chronological 1,727/750 train/test split:",
            48: "Baseline_LASSO: 27 market features; constant held-out prediction after regularisation",
            49: "Event_LASSO: all 92 features; identical constant prediction to the baseline",
            50: "XGBoost and LightGBM: 92-feature non-linear candidates; both overfit and fail promotion",
            52: "Target: next-day SPY log return. Baseline_LASSO test RMSE is 0.009631. Event_LASSO is identical; XGBoost (0.009656) and LightGBM (0.009700) are worse. No candidate passes the Bonferroni-corrected RMSE and directional-accuracy protocol.",
            55: "The study period is January 2015–December 2025. All core sources are free/public. The 4,018-row GDELT history is retained as a candidate contextual control, not current RQ evidence.",
            69: "Executed Python codebase: eight notebooks plus src/, tests, contracts and validation artefacts.",
            70: "Current event catalogue: 1,005 APP + FOMC events; 344 high-impact records; 264 usable CAR rows.",
            71: "Conditional pooled and category DoWhy estimates with explicit sparsity and identification caveats.",
            72: "FES v1.1 models, paired statistical tests, Random Forest importance and SHAP diagnostics.",
            73: "Synchronized dissertation draft, project summary, AI statement and Research Bible.",
            74: "Publication figures 08a–08d and six architecture SVGs, all validated/render-inspected.",
            105: "DATSCI7030 · Project Summary · Version 1.2 · FES v1.1 / MCP v1.0 · 15 July 2026",
        },
    )
    fill_table(doc.tables[0], [["Project Summary\nArchitecture · Methodology · Data Sources · Current Results\nVersion 1.2"]])
    fill_table(
        doc.tables[1],
        [
            ["Field", "Detail"],
            ["Author", "Ibrahim Haroun"],
            ["Module", "7030DATSCI — Data Science Project"],
            ["Institution", "Liverpool John Moores University (LJMU)"],
            ["Year", "2025–2026"],
            ["Version", "1.2 (FES v1.1 / MCP v1.0)"],
            ["Date Range", "January 2015–December 2025"],
        ],
    )
    fill_table(
        doc.tables[2],
        [["Core Research Questions\nRQ1: Do APP/FOMC communications produce abnormal SPY returns or a conditional sentiment effect?\nRQ2: Do event-derived features dominate predictive importance?\nRQ3: Do event-enhanced models beat a market-only baseline?"]],
    )
    fill_table(
        doc.tables[3],
        [["CAR Formula\nCARₑ(t₁,t₂) = Σ[Rₜ − μ̂], with current event window t₁=−5 and t₂=+10 trading days."]],
    )
    fill_table(
        doc.tables[4],
        [
            ["Source", "Type", "Access", "Current Volume"],
            ["Yahoo Finance", "SPY OHLCV and VIX", "Free", "2,765 SPY rows"],
            ["American Presidency Project", "Presidential documents", "Free", "10,892 raw; 916 analytical"],
            ["FRED", "Macroeconomic indicators", "Free", "Validated monthly series"],
            ["FOMC", "Meeting decisions", "Free", "89 events"],
            ["GDELT", "Candidate daily risk control", "Free", "4,018 rows; excluded from DAG/FES"],
        ],
    )
    fill_table(
        doc.tables[5],
        [["APP Scope\nAPP is the reproducible primary text archive. The current analytical subset contains 916 economic-policy titles; full-text scoring remains future work."]],
    )
    fill_table(
        doc.tables[6],
        [
            ["Category", "Tool / Library", "Purpose"],
            ["Language", "Python 3.13-compatible environment", "Pipeline implementation"],
            ["Data", "pandas, NumPy, PyArrow", "Wrangling and Parquet artefacts"],
            ["Statistics", "SciPy, statsmodels", "Event tests and diagnostics"],
            ["Causal inference", "DoWhy", "Backdoor estimate and refuters"],
            ["Machine learning", "scikit-learn, XGBoost, LightGBM", "Baseline and candidates"],
            ["NLP", "Transformers / FinBERT", "APP title sentiment"],
            ["Explainability", "Random Forest importance and SHAP", "RQ2 ranking and diagnostics"],
            ["Visualisation", "matplotlib, seaborn", "Validated publication figures"],
            ["Testing", "pytest plus validation JSON", "Code and artefact checks"],
            ["Governance", "Git and Research Bible", "Versioning and decision trail"],
        ],
    )
    fill_table(
        doc.tables[7],
        [
            ["Notebook", "Phase", "Current Status"],
            ["01_data_collection.ipynb", "Data ingestion", "Executed; raw sources refreshed"],
            ["02_eda.ipynb", "EDA / master dataset", "Executed; validation PASS"],
            ["03_event_detection.ipynb", "Event catalogue", "Executed; 1,005 events"],
            ["04_causal_analysis.ipynb", "CAR / DoWhy", "Executed; current outputs saved"],
            ["05_feature_engineering.ipynb", "FES v1.1", "Executed; validation PASS"],
            ["06_model_training.ipynb", "Market baseline", "Executed; validation PASS"],
            ["07_model_evaluation.ipynb", "RQ2/RQ3 evaluation", "Executed; validation PASS"],
            ["08_results_visualisation.ipynb", "Publication figures", "Executed; validation PASS"],
        ],
    )
    fill_table(
        doc.tables[8],
        [
            ["Risk", "Category", "Current Mitigation"],
            ["Title-only FinBERT domain mismatch", "NLP", "Report 92.4% neutral share; validate full text in future"],
            ["Clustered events", "Statistics", "BH-FDR is complete; disclose that it does not repair temporal dependence"],
            ["Sparse causal categories", "Causal inference", "Report n_nonzero and avoid strong category claims"],
            ["Tree-model overfitting", "Model", "Chronological holdout; report negative test R²"],
            ["Constant baseline directional accuracy", "Evaluation", "Treat as base-rate artefact; use RMSE/IC/paired tests"],
            ["GDELT design drift", "Governance", "Exclude from frozen DAG/FES; require versioned amendment"],
        ],
    )
    doc.core_properties.comments = "Synchronized with validated Notebooks 01–08 and the completed RQ1 reporting gate on 2026-07-15."
    doc.save(SUMMARY)


if __name__ == "__main__":
    sync_dissertation()
    sync_ai_statement()
    sync_project_summary()
    print("Synchronized:", MAIN, AI, SUMMARY, sep="\n- ")
